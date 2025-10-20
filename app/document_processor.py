import os
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv

load_dotenv()


class DocumentProcessor:
    """
    Handles document upload, text extraction, and chunking.
    Supports PDF and DOCX formats.
    """
    
    def __init__(self):
        """Initialize the document processor with text splitter"""
        # Configure the text splitter
        # This breaks documents into manageable chunks with overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,        # Characters per chunk
            chunk_overlap=100,      # Overlap between chunks (prevents losing context at boundaries)
            length_function=len,    # How to measure chunk size
            separators=[           # Try to split on these, in order of preference
                "\n\n",            # Paragraphs (best)
                "\n",              # Lines
                ".",               # Sentences
                " ",               # Words
                ""                 # Characters (last resort)
            ]
        )
        
        # Get upload directory from environment
        self.upload_dir = os.getenv("UPLOAD_DIR", "./uploads")
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_dir, exist_ok=True)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to the PDF file
        
        Returns:
            Extracted text as a single string
        """
        try:
            reader = PdfReader(file_path)
            text = ""
            
            # Extract text from each page
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num} ---\n{page_text}"
            
            print(f"✅ Extracted {len(text)} characters from PDF ({len(reader.pages)} pages)")
            return text
            
        except Exception as e:
            print(f"❌ Error extracting text from PDF: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to the DOCX file
        
        Returns:
            Extracted text as a single string
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from each paragraph
            for para_num, paragraph in enumerate(doc.paragraphs, start=1):
                if paragraph.text.strip():  # Skip empty paragraphs
                    text += paragraph.text + "\n"
            
            print(f"✅ Extracted {len(text)} characters from DOCX ({len(doc.paragraphs)} paragraphs)")
            return text
            
        except Exception as e:
            print(f"❌ Error extracting text from DOCX: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to the TXT file
        
        Returns:
            File contents as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            print(f"✅ Extracted {len(text)} characters from TXT")
            return text
            
        except Exception as e:
            print(f"❌ Error reading TXT file: {e}")
            raise
    
    def extract_text(self, file_path: str, filename: str) -> str:
        """
        Extract text from file based on extension
        
        Args:
            file_path: Path to the file
            filename: Original filename (to determine type)
        
        Returns:
            Extracted text
        """
        # Get file extension
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        elif extension == 'txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: .{extension}")
    
    def chunk_text(self, text: str, source: str) -> tuple[List[str], List[Dict[str, Any]]]:
        """
        Split text into chunks with metadata
        
        Args:
            text: The full text to chunk
            source: Source filename for metadata
        
        Returns:
            Tuple of (chunk_texts, chunk_metadatas)
        """
        # Split the text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create metadata for each chunk
        metadatas = []
        for i, chunk in enumerate(chunks):
            metadatas.append({
                "source": source,
                "chunk_id": i,
                "total_chunks": len(chunks),
                "char_count": len(chunk)
            })
        
        print(f"✅ Created {len(chunks)} chunks from {source}")
        return chunks, metadatas
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Complete pipeline: extract text, chunk it, and return results
        
        Args:
            file_path: Path to the uploaded file
            filename: Original filename
        
        Returns:
            Dict with chunks, metadatas, and stats
        """
        import time
        start_time = time.time()
        
        try:
            # Step 1: Extract text from file
            print(f"📄 Processing document: {filename}")
            text = self.extract_text(file_path, filename)
            
            # Check if text was extracted
            if not text or len(text.strip()) < 10:
                raise ValueError("No meaningful text extracted from document")
            
            # Step 2: Chunk the text
            chunks, metadatas = self.chunk_text(text, filename)
            
            # Step 3: Return results
            processing_time = time.time() - start_time
            return {
                "filename": filename,
                "chunks": chunks,
                "metadatas": metadatas,
                "total_chunks": len(chunks),
                "total_characters": len(text),
                "processing_time": processing_time,
                "status": "success"
            }
            
        except Exception as e:
            print(f"❌ Error processing document {filename}: {e}")
            return {
                "filename": filename,
                "status": "error",
                "error": str(e),
                "processing_time": time.time() - start_time
            }
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """
        Save uploaded file to disk
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
        
        Returns:
            Path to saved file
        """
        # Create safe filename (remove path traversal attempts)
        safe_filename = os.path.basename(filename)
        file_path = os.path.join(self.upload_dir, safe_filename)
        
        # Write file to disk
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        print(f"💾 Saved file to: {file_path}")
        return file_path
    
    def cleanup_file(self, file_path: str):
        """
        Delete uploaded file after processing
        
        Args:
            file_path: Path to file to delete
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"🗑️  Cleaned up: {file_path}")
        except Exception as e:
            print(f"⚠️  Could not delete {file_path}: {e}")


# Singleton instance
_document_processor_instance = None

def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor instance"""
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    return _document_processor_instance